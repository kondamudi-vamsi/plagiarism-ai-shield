import re
import math
import io
import hashlib
from collections import Counter
from typing import Dict, List, Tuple, Set, Any

# ==========================================
# 1. Document File Parsers (PDF & Word)
# ==========================================

def extract_text_from_pdf(content_bytes: bytes) -> str:
    """Extracts text from PDF byte stream."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content_bytes))
        text = []
        for page in reader.pages:
            val = page.extract_text()
            if val:
                text.append(val)
        return "\n".join(text)
    except Exception as e:
        return f"[PDF Extraction Error: {str(e)}]"

def extract_text_from_docx(content_bytes: bytes) -> str:
    """Extracts text from Word DOCX byte stream."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(content_bytes))
        text = []
        for p in doc.paragraphs:
            if p.text.strip():
                text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        return f"[Word DOCX Extraction Error: {str(e)}]"


# ==========================================
# 2. Text Plagiarism (TF-IDF & Cosine Similarity + Highlight)
# ==========================================

def clean_text_for_tfidf(text: str) -> List[str]:
    """Tokenizes text into words for vector similarity."""
    return re.findall(r'\b\w+\b', text.lower())

def get_text_similarity(text1: str, text2: str) -> Dict[str, Any]:
    """
    Computes overall Cosine Similarity score and locates
    exact/near matching sentence-level segments.
    """
    words1 = clean_text_for_tfidf(text1)
    words2 = clean_text_for_tfidf(text2)
    
    if not words1 or not words2:
        return {"score": 0.0, "matches": []}
    
    # Cosine Similarity
    c1 = Counter(words1)
    c2 = Counter(words2)
    all_words = set(c1.keys()).union(set(c2.keys()))
    dot_prod = sum(c1[w] * c2[w] for w in all_words)
    mag1 = math.sqrt(sum(v**2 for v in c1.values()))
    mag2 = math.sqrt(sum(v**2 for v in c2.values()))
    
    score = dot_prod / (mag1 * mag2) if (mag1 * mag2) > 0 else 0.0
    
    # Sentence-level comparison for match highlighting
    def get_sentences_with_offsets(text: str) -> List[Dict[str, Any]]:
        # Simple sentence splitter that tracks character bounds
        sentences = []
        # Find sentence boundaries
        pattern = re.compile(r'[^.!?\n\r]+[.!?\n\r]*')
        for match in pattern.finditer(text):
            s_text = match.group()
            s_clean = s_text.strip()
            # Filter out tiny sentences/whitespace to reduce noise
            if len(s_clean) > 15 and len(s_clean.split()) >= 3:
                sentences.append({
                    "text": s_clean,
                    "start": match.start(),
                    "end": match.end()
                })
        return sentences

    sents1 = get_sentences_with_offsets(text1)
    sents2 = get_sentences_with_offsets(text2)
    
    matches = []
    
    # Helpers for sentence similarity checking
    def normalize_sentence(s: str) -> Set[str]:
        return set(re.findall(r'\b\w+\b', s.lower()))

    # Track matched status to prevent highlighting same span multiple times
    matched2 = set()
    
    for i, s1 in enumerate(sents1):
        w1 = normalize_sentence(s1["text"])
        if not w1:
            continue
        
        best_match_idx = -1
        best_sim = 0.0
        
        for j, s2 in enumerate(sents2):
            if j in matched2:
                continue
            w2 = normalize_sentence(s2["text"])
            if not w2:
                continue
            
            # Jaccard index
            inter = len(w1.intersection(w2))
            union = len(w1.union(w2))
            sim = inter / union if union > 0 else 0.0
            
            if sim > best_sim:
                best_sim = sim
                best_match_idx = j
        
        # If similarity is high (>= 50%), consider it plagiary
        if best_sim >= 0.50 and best_match_idx != -1:
            matched2.add(best_match_idx)
            matches.append({
                "source": {
                    "text": s1["text"],
                    "start": s1["start"],
                    "end": s1["end"]
                },
                "target": {
                    "text": sents2[best_match_idx]["text"],
                    "start": sents2[best_match_idx]["start"],
                    "end": sents2[best_match_idx]["end"]
                },
                "similarity": best_sim
            })
            
    return {
        "score": float(score),
        "matches": matches
    }


# ==========================================
# 3. Code Plagiarism (Tokenization & Winnowing)
# ==========================================

# Standard keywords for major programming languages (Python, JS/TS, C++, Java, etc.)
CODE_KEYWORDS = {
    'def', 'class', 'import', 'from', 'as', 'return', 'if', 'elif', 'else', 'for', 'while', 'in', 'is', 'not',
    'and', 'or', 'try', 'except', 'finally', 'with', 'lambda', 'pass', 'break', 'continue', 'function', 'var',
    'let', 'const', 'new', 'this', 'typeof', 'instanceof', 'void', 'int', 'float', 'double', 'char', 'boolean',
    'bool', 'public', 'private', 'protected', 'static', 'final', 'extends', 'implements', 'interface', 'package',
    'switch', 'case', 'default', 'throw', 'throws', 'struct', 'namespace', 'include'
}

def tokenize_code(code: str) -> List[Dict[str, Any]]:
    """
    Tokenizes raw code into a stream of normalized token items.
    Filters out comments, whitespaces and categorizes everything else.
    """
    # Regex specifications
    rules = [
        ('COMMENT_MULTI_C', r'/\*[\s\S]*?\*/'),
        ('COMMENT_SINGLE_C', r'//.*'),
        ('COMMENT_PYTHON', r'#.*'),
        ('STRING_TRIPLE', r'"""[\s\S]*?"""|\'\'\'[\s\S]*?\'\'\''),
        ('STRING_SINGLE', r'"(?:\\.|[^"\\])*"|\'(?:\\.|[^\'\\])*\''),
        ('NUMBER', r'\b\d+(?:\.\d+)?\b'),
        ('IDENTIFIER', r'\b[a-zA-Z_][a-zA-Z0-9_]*\b'),
        ('SYMBOL', r'[+\-*/%&|^~<>!=]=?|[\(\)\[\]\{\};,.:\?]'),
        ('NEWLINE', r'\n'),
        ('SKIP', r'[ \t]'),
        ('MISMATCH', r'.'),
    ]
    
    master_re = re.compile('|'.join(f'(?P<{name}>{pattern})' for name, pattern in rules))
    tokens = []
    
    for match in master_re.finditer(code):
        token_type = match.lastgroup
        value = match.group()
        
        # Skip whitespaces and comments entirely to achieve structure invariance
        if token_type in ('SKIP', 'NEWLINE', 'COMMENT_MULTI_C', 'COMMENT_SINGLE_C', 'COMMENT_PYTHON'):
            continue
        
        # Categorize
        if token_type == 'STRING_TRIPLE' or token_type == 'STRING_SINGLE':
            # Normalize all string literals
            normalized_val = 'STR_LIT'
        elif token_type == 'NUMBER':
            normalized_val = 'NUM_LIT'
        elif token_type == 'IDENTIFIER':
            # Keep language keywords intact, collapse other variables/function names to general ID
            normalized_val = value if value in CODE_KEYWORDS else 'ID'
        elif token_type == 'SYMBOL':
            normalized_val = value
        else: # MISMATCH or unknown
            continue
            
        tokens.append({
            "type": token_type,
            "value": normalized_val,
            "orig": value,
            "start": match.start(),
            "end": match.end()
        })
        
    return tokens

def winnow(tokens: List[Dict[str, Any]], k: int = 8, w: int = 5) -> Dict[int, List[Tuple[int, int]]]:
    """
    Winnowing algorithm for generating search fingerprints.
    Returns: Dict mapping hash code -> list of (token_start_idx, token_end_idx)
    """
    if len(tokens) < k:
        # If too small, return hash of entire token list
        full_str = "".join(t["value"] for t in tokens)
        h = int(hashlib.md5(full_str.encode('utf-8')).hexdigest()[:8], 16)
        return {h: [(0, len(tokens) - 1)]} if tokens else {}

    # 1. Create k-grams and hash them
    kgrams = []
    hashes = []
    for i in range(len(tokens) - k + 1):
        kgram_tokens = tokens[i : i + k]
        kgram_string = "".join(t["value"] for t in kgram_tokens)
        
        # MD5 hash is stable across processes
        h = int(hashlib.md5(kgram_string.encode('utf-8')).hexdigest()[:8], 16)
        
        kgrams.append((i, i + k - 1)) # cover boundaries in token index
        hashes.append(h)
    
    # 2. Windowed selection of minimum hashes
    fingerprints = {} # hash -> list of token ranges
    
    if len(hashes) < w:
        # If hashes less than window size, pick the overall minimum
        min_idx = hashes.index(min(hashes))
        h = hashes[min_idx]
        fingerprints[h] = [kgrams[min_idx]]
        return fingerprints

    # Sliding window
    for i in range(len(hashes) - w + 1):
        window = hashes[i : i + w]
        # Find index of min value. In case of ties, pick the rightmost one (standard winnowing rule)
        min_val = min(window)
        # Scan backwards to get rightmost index in the window
        min_in_window_idx = -1
        for offset in range(w - 1, -1, -1):
            if window[offset] == min_val:
                min_in_window_idx = i + offset
                break
                
        h = hashes[min_in_window_idx]
        token_range = kgrams[min_in_window_idx]
        
        if h not in fingerprints:
            fingerprints[h] = []
        # Avoid duplicate range listings
        if token_range not in fingerprints[h]:
            fingerprints[h].append(token_range)
            
    return fingerprints

def get_code_similarity(code1: str, code2: str, k: int = 8, w: int = 5) -> Dict[str, Any]:
    """
    Performs Winnowing similarity checking on source code, returns matches and overall score.
    """
    tokens1 = tokenize_code(code1)
    tokens2 = tokenize_code(code2)
    
    if not tokens1 or not tokens2:
        return {"score": 0.0, "matches": []}
        
    fp1 = winnow(tokens1, k, w)
    fp2 = winnow(tokens2, k, w)
    
    hashes1 = set(fp1.keys())
    hashes2 = set(fp2.keys())
    
    shared_hashes = hashes1.intersection(hashes2)
    
    total_fps = len(hashes1) + len(hashes2)
    score = (2.0 * len(shared_hashes)) / total_fps if total_fps > 0 else 0.0
    
    # Locate highlight ranges in original source texts
    matches = []
    
    for h in shared_hashes:
        # Maps matching fingerprint hashes back to char offsets in both code files
        ranges1 = fp1[h]
        ranges2 = fp2[h]
        
        for r1 in ranges1:
            start_t1, end_t1 = r1
            start_char1 = tokens1[start_t1]["start"]
            end_char1 = tokens1[end_t1]["end"]
            chunk1 = code1[start_char1:end_char1]
            
            for r2 in ranges2:
                start_t2, end_t2 = r2
                start_char2 = tokens2[start_t2]["start"]
                end_char2 = tokens2[end_t2]["end"]
                chunk2 = code2[start_char2:end_char2]
                
                matches.append({
                    "hash": h,
                    "source": {
                        "text": chunk1,
                        "start": start_char1,
                        "end": end_char1
                    },
                    "target": {
                        "text": chunk2,
                        "start": start_char2,
                        "end": end_char2
                    }
                })
                
    # Sort matches by source start position to simplify rendering
    matches.sort(key=lambda x: x["source"]["start"])
    
    return {
        "score": float(score),
        "matches": matches
    }
